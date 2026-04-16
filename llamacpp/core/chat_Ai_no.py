#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
轻舟 AI・LightShip AI - 聊天 (llama.cpp 后端) + 在线AI入口
适配 Gradio 6.2.0，支持多模态看图模型 (qwen3.5, gemma4 等)
Copyright 2026 光影的故事2018
"""

import sys
import os
SCRIPT_DIR = os.path.dirname(__file__)
sys.path.insert(0, SCRIPT_DIR)

import gradio as gr
import requests
import json
import time
import re
import base64
import webbrowser
import tempfile
import threading
import imghdr
from docx import Document
from datetime import datetime

# ==================== llama.cpp 配置 ====================
LLAMA_API_URL = "http://127.0.0.1:8080/v1/chat/completions"
MODELS_URL = "http://127.0.0.1:8080/v1/models"

# ==================== 全局停止标志（基于会话） ====================
_stop_flags = {}
_stop_lock = threading.Lock()

# ==================== llama.cpp 服务与模型检查 ====================
def is_llama_available():
    try:
        resp = requests.get(MODELS_URL, timeout=3)
        return resp.status_code == 200, "服务正常" if resp.status_code == 200 else f"服务异常，状态码 {resp.status_code}"
    except Exception as e:
        return False, f"服务未启动或连接失败：{str(e)}"

def get_llama_models():
    try:
        resp = requests.get(MODELS_URL, timeout=5)
        if resp.status_code == 200:
            data = resp.json()
            models = [item["id"] for item in data.get("data", [])]
            return models
    except Exception as e:
        print(f"获取模型列表失败: {e}")
    return []

def is_multimodal(model_name: str) -> bool:
    multimodal_keywords = ["qwen", "llava", "bakllava", "gemini", "cogvlm", "minicpm", "gemma"]
    return any(kw in model_name.lower() for kw in multimodal_keywords)

def get_model_display_list(models):
    display_list = []
    for m in models:
        if is_multimodal(m):
            display_list.append((f"{m} (多模态)", m))
        elif "deepseek" in m.lower() or "r1" in m.lower():
            display_list.append((f"{m} (深度推理)", m))
        else:
            display_list.append((m, m))
    return display_list

# ==================== 在线AI网站 URL ====================
URLS = {
    "DeepL": "https://www.deepl.com/zh",
    "有道翻译": "https://fanyi.youdao.com/",
    "豆包": "https://www.doubao.com/",
    "通义千问": "https://www.qianwen.com/",
    "DeepSeek": "https://www.deepseek.com/",
    "ChatGLM": "https://chatglm.cn",
    "Kimi": "https://kimi.moonshot.cn/",
    "腾讯元宝": "https://yuanbao.tencent.com/",
}

# ==================== 预置提示词 ====================
PROMPTS = {
    "SRT字幕翻译(保留时间码)": """你是一个专业的字幕翻译专家。请将以下SRT字幕内容翻译成中文。
要求：
1. 严格保留原文的时间轴格式 (例如: 00:00:01,000 --> 00:00:03,000)。
2. 保持字幕序号不变。
3. 翻译要自然流畅，适合口语表达，注意联系上下文语境。
4. 对于专有名词或专业术语，请保持一致性。
5. 直接输出翻译后的SRT内容，不要包含解释。

原文内容：
""",
    "语义断句与合并(ASR优化)": """你是专业字幕精修师。
我将给你ASR识别的带时间戳字幕，可能是碎句、断句错误、错别字。
请按以下规则处理：

1. 按语义合并碎句，不要在一句话中间切开。
2. 每条字幕中文字数 ≤ 20 字。
3. 合并多条时，开始时间=第一条，结束时间=最后一条。
4. 自动修正ASR错别字、同音字错误。
5. 去掉口语冗余词：那个、就是、然后、嘛、啊等。
6. 输出严格标准SRT格式，不要任何解释、不要JSON。

输入内容：
""",
    "短视频极速版(10字以内)": """你是短视频字幕专家。
规则：
1. 每一条字幕 ≤ 12 个字。
2. 必须按语义断句,不许硬切词、不许乱时序。
3. 短句、有力、适合短视频节奏。
4. 自动合并、自动拆分过长句。
5. 保留时间码，输出标准SRT。

输入内容：
""",
    "双语字幕版(中英对照)": """你是专业字幕翻译。
规则：
1. 保留原时间轴。
2. 中文在上，英文在下。
3. 语言自然口语化，不生硬。
4. 输出标准SRT。

输入内容：
""",
    "双语字幕生成(中英对照)": """请将以下文本翻译成中文，并生成中英双语对照格式。
格式要求：
第一行：原文
第二行：译文
(空行分隔不同段落)

请处理以下内容：
""",
    "文本润色与校对": """请对以下文本进行润色和校对。
要求：
1. 修正错别字和标点符号错误。
2. 优化语句通顺度，使其更符合中文阅读习惯。
3. 保持原意不变，不增加或删减关键信息。
4. 如果没有错误，请原样输出。

待处理文本：
""",
    "长文本总结(适合Kimi/元宝)": """请阅读以下长文本，并进行总结。
要求：
1. 提炼出核心观点和关键信息。
2. 使用条理清晰的列表形式输出。
3. 语言简洁明了。

文本内容：
""",
    "ASR校对与纠错(专业版)": """你是一个专业的校对员。以下是 ASR 语音识别的原始结果，可能包含错别字。
请根据上下文修正错别字，并输出修正后的纯文本内容。

要求：
1. 重点修正同音字错误。
2. 保持原意，不要大幅改写。
3. 输出修正后的文本即可。

原文内容：
"""
}

def open_url(url):
    webbrowser.open(url)
    return f"✅ 已打开 {url}，请查看浏览器"

def update_prompt(prompt_name):
    return PROMPTS.get(prompt_name, "")

# ==================== 图片编码（用于多模态） ====================
def encode_image_to_base64(image_path):
    try:
        img_format = imghdr.what(image_path)
        if img_format is None:
            ext = os.path.splitext(image_path)[1].lower()
            if ext in ['.jpg', '.jpeg']:
                img_format = 'jpeg'
            elif ext == '.png':
                img_format = 'png'
            elif ext == '.gif':
                img_format = 'gif'
            else:
                img_format = 'jpeg'
        mime_type = f"image/{img_format}"
        with open(image_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        return f"data:{mime_type};base64,{b64}"
    except Exception as e:
        print(f"图片编码失败: {e}")
        return None

# ==================== 流式解析器（修复乱码：UTF-8 安全清洗）====================
class StreamResponseParser:
    def __init__(self):
        self.thought = ""
        self.answer = ""
        self.buffer = ""
        self.in_think_tag = False
        self.has_think = False
        self.start_time = time.time()
        self.total_tokens = 0
        self.char_count = 0

    def _clean_text(self, text: str) -> str:
        """移除或替换无效的 UTF-8 字符，防止乱码"""
        if not text:
            return text
        return text.encode('utf-8', errors='replace').decode('utf-8')

    def parse_chunk(self, chunk_data: dict) -> dict:
        result = {"thought": "", "answer": "", "status": "answering"}
        delta = chunk_data.get("choices", [{}])[0].get("delta", {})
        
        reasoning = delta.get("reasoning_content", "")
        if reasoning:
            clean_reasoning = self._clean_text(reasoning)
            self.thought += clean_reasoning
            self.has_think = True
            result["thought"] = clean_reasoning
            result["status"] = "thinking"
            return result
        
        content = delta.get("content", "")
        if not content:
            return result
        
        content = self._clean_text(content)
        self.char_count += len(content)
        self.buffer += content
        
        while True:
            if not self.in_think_tag:
                start_idx = self.buffer.find('<think>')
                if start_idx != -1:
                    before_think = self.buffer[:start_idx]
                    if before_think:
                        before_think = self._clean_text(before_think)
                        self.answer += before_think
                        result["answer"] += before_think
                    self.buffer = self.buffer[start_idx + 7:]
                    self.in_think_tag = True
                    self.has_think = True
                    result["status"] = "thinking"
                    continue
                else:
                    if self.buffer.rstrip().endswith('<') or '<' in self.buffer[-5:]:
                        last_lt = self.buffer.rfind('<')
                        if last_lt != -1:
                            safe_part = self.buffer[:last_lt]
                            self.buffer = self.buffer[last_lt:]
                            if safe_part:
                                safe_part = self._clean_text(safe_part)
                                self.answer += safe_part
                                result["answer"] += safe_part
                        else:
                            clean_buf = self._clean_text(self.buffer)
                            self.answer += clean_buf
                            result["answer"] += clean_buf
                            self.buffer = ""
                    else:
                        clean_buf = self._clean_text(self.buffer)
                        self.answer += clean_buf
                        result["answer"] += clean_buf
                        self.buffer = ""
                    break
            else:
                end_idx = self.buffer.find('</think>')
                if end_idx != -1:
                    think_part = self.buffer[:end_idx]
                    if think_part:
                        think_part = self._clean_text(think_part)
                        self.thought += think_part
                        result["thought"] += think_part
                    self.buffer = self.buffer[end_idx + 8:]
                    self.in_think_tag = False
                    result["status"] = "answering"
                    continue
                else:
                    clean_buf = self._clean_text(self.buffer)
                    self.thought += clean_buf
                    result["thought"] += clean_buf
                    self.buffer = ""
                    break
        return result

    def finalize(self, usage=None):
        if self.buffer:
            clean_buffer = self._clean_text(self.buffer)
            if self.in_think_tag:
                self.thought += clean_buffer
            else:
                self.answer += clean_buffer
            self.buffer = ""
        
        if not self.has_think and self.thought:
            self.answer = self.thought + self.answer
            self.thought = ""
        
        if usage:
            self.total_tokens = usage.get("completion_tokens", 0)
        else:
            self.total_tokens = self.char_count // 2
        
        return self.answer, self.thought

    def get_processing_time(self):
        duration = time.time() - self.start_time
        if duration < 60:
            return f"耗时：{duration:.1f}秒"
        else:
            minutes = int(duration // 60)
            seconds = duration % 60
            return f"耗时：{minutes}分{seconds:.1f}秒"

# ==================== 聊天核心函数（llama.cpp 后端）====================
def build_messages(message, image_data_url, model_name):
    system_content = f"今天是 {datetime.now().strftime('%Y年%m月%d日')}。你是一个智能助手，请用中文回答。"
    messages = [{"role": "system", "content": system_content}]
    if image_data_url and is_multimodal(model_name):
        user_content = [
            {"type": "text", "text": message},
            {"type": "image_url", "image_url": {"url": image_data_url}}
        ]
        messages.append({"role": "user", "content": user_content})
    else:
        messages.append({"role": "user", "content": message})
    return messages

def stream_response_llama(message, image_path, model_name, temperature, max_tokens, gpu_layers, history, request: gr.Request):
    global _stop_flags, _stop_lock
    
    available, status_msg = is_llama_available()
    if not available:
        error_content = f"❌ {status_msg}，请先启动 llama-server。"
        updated_history = history + [{"role": "assistant", "content": error_content}]
        yield updated_history, status_msg
        return
    
    models = get_llama_models()
    if model_name not in models:
        error_content = f"❌ 模型 {model_name} 不在 llama-server 可用列表中。请检查模型名称或刷新列表。"
        updated_history = history + [{"role": "assistant", "content": error_content}]
        yield updated_history, "模型不可用"
        return
    
    session_id = request.session_hash
    with _stop_lock:
        _stop_flags[session_id] = False
    
    image_data_url = None
    if image_path is not None:
        if is_multimodal(model_name):
            image_data_url = encode_image_to_base64(image_path)
            if not image_data_url:
                image_path = None
        else:
            print(f"警告：模型 {model_name} 不支持图片，已忽略。")
            image_path = None
    
    messages = build_messages(message, image_data_url, model_name)
    # 基础 payload
    payload = {
        "model": model_name,
        "messages": messages,
        "stream": True,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    # 只有 gpu_layers != -1 时才添加 GPU 层数字段
    if gpu_layers != -1:
        payload["ngl"] = gpu_layers
        payload["n_gpu_layers"] = gpu_layers
    
    timestamp = time.strftime('%H:%M:%S')
    user_content = f"[{timestamp}] 用户：{message}" + (" [附图片]" if image_path else "")
    updated_history = history + [{"role": "user", "content": user_content}]
    updated_history.append({"role": "assistant", "content": f"[{timestamp}] {model_name}："})
    gpu_display = "服务器默认" if gpu_layers == -1 else str(gpu_layers)
    yield updated_history, f"模型 [{model_name}] 正在生成... (GPU层数: {gpu_display})"
    
    parser = StreamResponseParser()
    full_answer = ""
    full_thought = ""
    
    try:
        response = requests.post(LLAMA_API_URL, json=payload, stream=True, timeout=180)
        response.raise_for_status()
        response.encoding = 'utf-8'
        
        for line in response.iter_lines(decode_unicode=True):
            with _stop_lock:
                if _stop_flags.get(session_id, False):
                    updated_history[-1]["content"] = f"[{timestamp}] {model_name}：{full_answer}\n\n[已手动停止]"
                    yield updated_history, "生成已停止"
                    return
            
            if line:
                line = line.strip()
                if line.startswith("data: "):
                    data_str = line[6:]
                    if data_str == "[DONE]":
                        continue
                    try:
                        chunk = json.loads(data_str)
                        parsed = parser.parse_chunk(chunk)
                        if parsed["thought"]:
                            full_thought += parsed["thought"]
                        if parsed["answer"]:
                            full_answer += parsed["answer"]
                        
                        display = f"[{timestamp}] {model_name}："
                        if full_thought:
                            display += f"\n\n💭 {full_thought}\n\n"
                        if full_answer:
                            display += full_answer
                        else:
                            display += "█"
                        updated_history[-1]["content"] = display
                        yield updated_history, f"生成中... ({len(full_answer)} 字符)"
                        
                        if "usage" in chunk:
                            parser.total_tokens = chunk["usage"].get("completion_tokens", 0)
                    except json.JSONDecodeError:
                        continue
        
        end_time = time.time()
        time_cost = end_time - parser.start_time
        final_answer, final_thought = parser.finalize()
        total_tokens = parser.total_tokens if parser.total_tokens > 0 else parser.char_count // 2
        speed = total_tokens / time_cost if time_cost > 0 else 0
        stat_str = f"{total_tokens} tokens, {time_cost:.1f}s, {speed:.2f}t/s"
        
        thought_html = ""
        if final_thought:
            safe_thought = final_thought.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            thought_html = f'''
            <details class="thoughts-details">
                <summary><strong>💭 思考过程</strong> (点击展开/折叠)</summary>
                <div class="thoughts-content"><em>{safe_thought.replace(chr(10), '<br>')}</em></div>
            </details>
            '''
        final_content = f"[{timestamp}] {model_name}："
        if thought_html:
            final_content += f"\n\n{thought_html}\n\n"
        final_content += f"{final_answer}\n\n[统计] {stat_str} (GPU层数: {gpu_display})"
        updated_history[-1]["content"] = final_content
        yield updated_history, f"生成完成，{stat_str}"
        
    except requests.exceptions.ConnectionError:
        updated_history[-1]["content"] = f"[{timestamp}] {model_name}：连接失败，请确保已运行 llama-server"
        yield updated_history, "连接失败"
    except Exception as e:
        updated_history[-1]["content"] = f"[{timestamp}] {model_name}：错误：{str(e)}"
        yield updated_history, f"错误：{str(e)}"
    finally:
        with _stop_lock:
            _stop_flags.pop(session_id, None)

def stop_generation(request: gr.Request):
    session_id = request.session_hash
    with _stop_lock:
        _stop_flags[session_id] = True
    return "正在停止生成..."

def clear_all():
    initial_history = [{
        "role": "assistant",
        "content": """AI本地小助手 (llama.cpp 后端)

使用指南：
1. 选择支持多模态的模型（如 qwen3.5 系列、gemma4 系列）以使用图片识别
2. 温度控制创意度，最大长度控制回复长度
3. 对话支持思考过程显示（如 DeepSeek-R1 风格）
4. 可随时停止生成
5. 调整 GPU 层数（-1=服务器默认，0=仅CPU，1~99=指定层数）

开始对话："""
    }]
    return initial_history, "对话已清空"

def check_llama_status():
    available, msg = is_llama_available()
    if available:
        models = get_llama_models()
        if models:
            return f"✅ llama.cpp 服务正常，可用模型：{', '.join(models[:5])}{'...' if len(models)>5 else ''}"
        else:
            return "✅ llama.cpp 服务正常，但未检测到任何模型"
    else:
        return f"❌ {msg}"

def refresh_models():
    models = get_llama_models()
    if not models:
        return gr.Dropdown(choices=[("请先启动 llama-server", "none")], value="none")
    choices = get_model_display_list(models)
    return gr.Dropdown(choices=choices, value=choices[0][1] if choices else None)

# ==================== 导出对话功能 ====================
def strip_html_tags(text) -> str:
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    text = re.sub(r'<details[^>]*>', '', text)
    text = re.sub(r'</details>', '', text)
    text = re.sub(r'<summary[^>]*>', '', text)
    text = re.sub(r'</summary>', '', text)
    text = re.sub(r'<div[^>]*>', '', text)
    text = re.sub(r'</div>', '', text)
    text = re.sub(r'<em>', '', text)
    text = re.sub(r'</em>', '', text)
    text = re.sub(r'<strong>', '', text)
    text = re.sub(r'</strong>', '', text)
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def export_chat_txt(history):
    if not history:
        return None
    lines = []
    for msg in history:
        role = msg.get("role", "unknown")
        content = msg.get("content")
        clean_content = strip_html_tags(content)
        if not clean_content:
            continue
        if role == "user":
            lines.append(f"【用户】\n{clean_content}\n")
        else:
            lines.append(f"【助手】\n{clean_content}\n")
        lines.append("-" * 50 + "\n")
    if not lines:
        return None
    full_text = "\n".join(lines)
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write(full_text)
        return f.name

def export_chat_docx(history):
    if not history:
        return None
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    title = doc.add_heading("轻舟 AI 对话记录", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for msg in history:
        role = msg.get("role", "unknown")
        content = msg.get("content")
        clean_content = strip_html_tags(content)
        if not clean_content:
            continue
        
        p = doc.add_paragraph()
        run = p.add_run(f"【{role}】\n")
        run.bold = True if role == "assistant" else False
        run.font.size = Pt(11)
        
        p2 = doc.add_paragraph(clean_content)
        p2.paragraph_format.left_indent = Pt(20)
        p2.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph("_" * 50)
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name

# ==================== Gradio 界面 ====================
logo_path = os.path.join(SCRIPT_DIR, "ai_logo.png")
ico_path = os.path.join(SCRIPT_DIR, "ai_logo.ico")

initial_models = get_llama_models()
if not initial_models:
    initial_models = ["请先启动 llama-server"]
model_choices = get_model_display_list(initial_models)

css = """
.thoughts-details {
    border: 1px solid var(--border-color-primary);
    border-radius: 6px;
    padding: 10px;
    margin: 10px 0;
    background-color: var(--color-background-tertiary);
}
.thoughts-details summary {
    cursor: pointer;
    padding: 8px;
    font-weight: bold;
    border-radius: 4px;
    background-color: var(--color-background-primary);
}
.thoughts-content {
    padding: 10px;
    background-color: var(--color-background-secondary);
    border-radius: 4px;
    font-style: italic;
}
"""

with gr.Blocks(title="轻舟 AI・LightShip AI (llama.cpp 后端)") as demo:
    with gr.Row():
        if os.path.exists(logo_path):
            gr.Image(logo_path, height=50, show_label=False, container=False, scale=0)
        with gr.Column(scale=1):
            gr.Markdown("""
            # 轻舟 AI・LightShip AI (llama.cpp)
            ###### 轻舟渡万境，一智载千寻。
            """)
    
    with gr.Tabs():
        # ---------- Tab1: 聊天（llama.cpp）----------
        with gr.Tab("聊天"):
            with gr.Row():
                with gr.Column(scale=4):
                    history_box = gr.Chatbot(
                        value=[{
                            "role": "assistant",
                            "content": """AI本地小助手 (llama.cpp 后端)

使用指南：
1. 选择支持多模态的模型（如 qwen3.5 系列、gemma4 系列）以使用图片识别
2. 温度控制创意度，最大长度控制回复长度
3. 对话支持思考过程显示（如 DeepSeek-R1 风格）
4. 可随时停止生成
5. 调整 GPU 层数（-1=服务器默认，0=仅CPU，1~99=指定层数）

开始对话："""
                        }],
                        height=900,
                        sanitize_html=False
                    )
                    with gr.Row():
                        input_box = gr.Textbox(
                            label="输入文字",
                            placeholder="请输入问题... (按回车发送)",
                            lines=3,
                            scale=4
                        )
                        send_btn = gr.Button("发送", variant="primary", scale=1)
                        stop_btn = gr.Button("停止", variant="stop", scale=1)
                
                with gr.Column(scale=1, min_width=250):
                    gr.Markdown("### 模型设置")
                    model_select = gr.Dropdown(
                        choices=model_choices,
                        value=model_choices[0][1] if model_choices else None,
                        label="选择模型"
                    )
                    refresh_btn = gr.Button("🔄 刷新模型列表", size="sm")
                    temp_slider = gr.Slider(0.1, 2.0, value=0.7, step=0.1, label="温度")
                    token_slider = gr.Slider(512, 8192, value=2048, step=128, label="最大长度")
                    # 修改：范围 -1 ~ 99，默认 -1
                    gpu_layers_slider = gr.Slider(
                        -1, 99, value=-1, step=1,
                        label="GPU 层数 (-1=服务器默认, 0=仅CPU, 1~99=指定层数)",
                        info="需 llama-server 支持请求级设置 (b8690 已支持)"
                    )
                    gr.Markdown("---")
                    gr.Markdown("### 图片上传")
                    image_input = gr.Image(
                        label="仅多模态模型支持（qwen3.5, gemma4 等）",
                        type="filepath",
                        height=180
                    )
                    gr.Markdown("---")
                    check_btn = gr.Button("检查服务", variant="secondary")
                    clear_btn = gr.Button("清空对话", variant="secondary")
                    # 修改：添加 lines=2
                    status_box = gr.Textbox(label="状态", value="就绪", interactive=False, lines=2)
                    
                    gr.Markdown("### 导出对话")
                    with gr.Row():
                        export_txt_btn = gr.Button("📄 导出为 TXT", variant="secondary", size="sm")
                        export_docx_btn = gr.Button("📝 导出为 Word", variant="secondary", size="sm")
                    export_file = gr.File(label="下载文件", visible=False)
            
            refresh_btn.click(fn=refresh_models, outputs=[model_select])
            send_btn.click(
                fn=stream_response_llama,
                inputs=[input_box, image_input, model_select, temp_slider, token_slider, gpu_layers_slider, history_box],
                outputs=[history_box, status_box]
            ).then(lambda: ("", None), None, [input_box, image_input])
            
            input_box.submit(
                fn=stream_response_llama,
                inputs=[input_box, image_input, model_select, temp_slider, token_slider, gpu_layers_slider, history_box],
                outputs=[history_box, status_box]
            ).then(lambda: ("", None), None, [input_box, image_input])
            
            stop_btn.click(fn=stop_generation, outputs=[status_box])
            clear_btn.click(fn=clear_all, outputs=[history_box, status_box])
            check_btn.click(fn=check_llama_status, outputs=[status_box])
            demo.load(fn=check_llama_status, outputs=[status_box])
            
            export_txt_btn.click(
                fn=export_chat_txt,
                inputs=[history_box],
                outputs=[export_file]
            ).then(lambda: gr.update(visible=True), None, [export_file])
            
            export_docx_btn.click(
                fn=export_chat_docx,
                inputs=[history_box],
                outputs=[export_file]
            ).then(lambda: gr.update(visible=True), None, [export_file])
        
        # ---------- Tab2: 在线AI入口 ----------
        with gr.Tab("在线AI入口"):
            with gr.Column():
                gr.Markdown("### 快速入口")
                with gr.Row(equal_height=True):
                    btn_deepl = gr.Button("DeepL", variant="secondary")
                    btn_youdao = gr.Button("有道翻译", variant="secondary")
                    btn_deepseek = gr.Button("DeepSeek", variant="secondary")
                    btn_doubao = gr.Button("豆包", variant="secondary")
                with gr.Row(equal_height=True):
                    btn_qianwen = gr.Button("通义千问", variant="secondary")
                    btn_kimi = gr.Button("Kimi", variant="secondary")
                    btn_chatglm = gr.Button("ChatGLM", variant="secondary")
                    btn_yuanbao = gr.Button("腾讯元宝", variant="secondary")
                status_trans = gr.Textbox(label="", value="等待操作...", interactive=False)
                
                gr.Markdown("---")
                gr.Markdown("### 提示词模板")
                prompt_selector = gr.Dropdown(
                    label="选择提示词类型",
                    choices=list(PROMPTS.keys()),
                    value="SRT字幕翻译(保留时间码)",
                )
                prompt_display = gr.Textbox(
                    label="提示词内容 (可直接编辑)",
                    value=PROMPTS["SRT字幕翻译(保留时间码)"],
                    lines=10,
                    interactive=True,
                )
                gr.Markdown("💡 **使用方法**：选择模板 → 复制提示词 → 点击上方按钮打开网站 → 粘贴使用。")
                
                btn_deepl.click(fn=lambda: open_url(URLS["DeepL"]), outputs=status_trans)
                btn_youdao.click(fn=lambda: open_url(URLS["有道翻译"]), outputs=status_trans)
                btn_deepseek.click(fn=lambda: open_url(URLS["DeepSeek"]), outputs=status_trans)
                btn_doubao.click(fn=lambda: open_url(URLS["豆包"]), outputs=status_trans)
                btn_qianwen.click(fn=lambda: open_url(URLS["通义千问"]), outputs=status_trans)
                btn_kimi.click(fn=lambda: open_url(URLS["Kimi"]), outputs=status_trans)
                btn_chatglm.click(fn=lambda: open_url(URLS["ChatGLM"]), outputs=status_trans)
                btn_yuanbao.click(fn=lambda: open_url(URLS["腾讯元宝"]), outputs=status_trans)
                prompt_selector.change(fn=update_prompt, inputs=[prompt_selector], outputs=[prompt_display])
    
    # 页脚
    gr.Markdown("---")
    gr.HTML("""
    <div style="text-align: center; color: #666; font-size: 0.9em;">
        <p>本工具仅用于个人学习与视频剪辑使用，禁止商业用途。</p>
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 8px; margin: 15px auto; max-width: 600px;">
            <p style="color: white; font-weight: bold; margin: 5px 0;">🎬 更新请关注B站up主：光影的故事2018</p>
            <p style="color: white; margin: 5px 0;">
                🔗 <strong>B站主页</strong>: <a href="https://space.bilibili.com/381518712" target="_blank" style="color: #ffdd40; text-decoration: none;">space.bilibili.com/381518712</a>
            </p>
        </div>
        <p>© 原创 WebUI 代码 © 2026 光影紐扣 版权所有</p>
    </div>
    """)

# ==================== 启动 ====================
if __name__ == "__main__":
    print("=" * 60)
    print("启动 轻舟 AI・LightShip AI (llama.cpp 后端)")
    print("聊天请确保已运行: llama-server.exe -m model.gguf --host 0.0.0.0 --port 8080")
    print("=" * 60)
    
    status = check_llama_status()
    print(status)
    if "未启动" in status:
        print("\n警告：llama.cpp 服务未启动，聊天功能将无法使用。")
    
    ports_to_try = [7866, 7961, 7861, 7862, 7960]
    for port in ports_to_try:
        try:
            demo.launch(
                server_name="127.0.0.1",
                server_port=port,
                share=False,
                inbrowser=True,
                theme=gr.themes.Default(),
                css=css,
                favicon_path=ico_path if os.path.exists(ico_path) else None
            )
            break
        except OSError as e:
            print(f"端口 {port} 启动失败: {e}")
            continue
    else:
        print("所有尝试的端口均被占用或不可用，请手动指定空闲端口。")